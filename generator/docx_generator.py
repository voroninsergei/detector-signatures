"""
Functions for generating expert conclusion reports in DOCX format.

Stage 7 of the development plan calls for automatic generation of a
forensic handwriting expert report.  This module provides a simple
interface to build a Word document containing the following sections:

* Introductory information (purpose of the study, received samples,
  sources of comparison).
* Methodology and course of the research (description of applied
  methods, preprocessing steps, feature extraction, diagnostic and
  comparative analyses).
* Results of the comparative study (tables of matching/non‑matching
  features, similarity scores, graphical material).
* Diagnostic conclusions (probabilistic statements about gender, age
  or other characteristics).
* Final conclusion and expert statement.
* Attachments (illustrations of selected fragments, charts, etc.).

The primary entry point is :func:`generate_conclusion`, which accepts
structured data produced by earlier stages and writes a DOCX file.

This implementation relies on the ``python‑docx`` package.  If the
package is not installed, install it via pip using the command
``pip install python-docx``.
"""

from typing import Dict, List, Any, Optional
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
except ImportError as exc:  # pragma: no cover
    raise RuntimeError(
        "python-docx is required to use the report generator. "
        "Install it with 'pip install python-docx'."
    ) from exc


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    """Add a heading to the document with a consistent style."""
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(14 if level == 1 else 12)
    return None


def _add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    """Add a normal paragraph to the document with optional bold text."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    run.bold = bold
    return None


def _add_image(document: Document, image_path: str, width_in_inches: float = 5.0) -> None:
    """Add an image to the document if the file exists."""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=Inches(width_in_inches))


def generate_conclusion(data: Dict[str, Any], output_path: str) -> None:
    """
    Generate a DOCX expert conclusion document.

    Parameters
    ----------
    data : dict
        Structured data containing the components of the report.  Expected
        keys include:

        * ``intro`` (str) – introductory information about the case.
        * ``methodology`` (str) – description of the applied methods.
        * ``comparative_analysis`` (str) – narrative of comparative results.
        * ``diagnostic`` (str) – narrative of diagnostic findings.
        * ``conclusion`` (str) – final expert statement.
        * ``images`` (list[str]) – list of file paths to images (optional).
        * ``tables`` (list[list[list[str]]]) – list of tables; each table
          is a list of rows, each row is a list of strings.

        Any missing sections will be omitted.

    output_path : str
        The path where the generated DOCX file will be saved.

    Returns
    -------
    None
    """
    document = Document()

    # Title
    _add_heading(document, 'Экспертное заключение', level=1)
    _add_paragraph(document, f"Дата: {datetime.now().strftime('%d.%m.%Y')}")

    # Introductory information
    intro = data.get('intro')
    if intro:
        _add_heading(document, '1. Вводные данные', level=2)
        _add_paragraph(document, intro)

    # Methodology and course of research
    methodology = data.get('methodology')
    if methodology:
        _add_heading(document, '2. Методика и ход исследования', level=2)
        _add_paragraph(document, methodology)

    # Comparative analysis section
    comp_analysis = data.get('comparative_analysis')
    if comp_analysis:
        _add_heading(document, '3. Результаты сравнительного исследования', level=2)
        _add_paragraph(document, comp_analysis)

    # Diagnostic section
    diagnostic = data.get('diagnostic')
    if diagnostic:
        _add_heading(document, '4. Диагностические выводы', level=2)
        _add_paragraph(document, diagnostic)

    # Insert tables (if any)
    tables = data.get('tables', [])
    for table_data in tables:
        table = document.add_table(rows=1, cols=len(table_data[0]))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_data[0]):
            hdr_cells[i].text = header
        for row_data in table_data[1:]:
            row_cells = table.add_row().cells
            for j, cell_val in enumerate(row_data):
                row_cells[j].text = cell_val
        document.add_paragraph()  # spacing after table

    # Images
    images: Optional[List[str]] = data.get('images')
    if images:
        _add_heading(document, '5. Иллюстрации', level=2)
        for img_path in images:
            _add_image(document, img_path, width_in_inches=5.0)
            document.add_paragraph()

    # Final conclusion
    conclusion = data.get('conclusion')
    if conclusion:
        _add_heading(document, '6. Заключение', level=2)
        _add_paragraph(document, conclusion, bold=True)

    # Save the document
    document.save(output_path)
